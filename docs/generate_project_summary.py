from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\Users\Naresh Suthar\vscode projects\supply\supplychain-auth\docs\SUPPLYPRINT_PROJECT_SUMMARY.docx"
BLUE = "2E74B5"; DARK = "1F4D78"; MUTED = "617187"; PALE = "E8EEF5"

def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"; run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, value):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), value); tc_pr.append(shd)

def set_cell(cell, text, bold=False):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); r = p.add_run(text); font(r, 9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True); shade(t.rows[0].cells[i], PALE)
        if widths: t.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], str(value))
            if widths: cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def heading(doc, text, level=1):
    p = doc.add_paragraph(); p.style = f"Heading {level}"; p.paragraph_format.keep_with_next = True
    r = p.add_run(text); font(r, 16 if level == 1 else 13 if level == 2 else 12, True, BLUE if level < 3 else DARK)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.15; font(p.add_run(text), 11)

def para(doc, text, bold_lead=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.1
    if bold_lead:
        font(p.add_run(bold_lead), 11, True); font(p.add_run(text), 11)
    else: font(p.add_run(text), 11)
    return p

doc = Document(); sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)
styles = doc.styles
styles["Normal"].font.name = "Calibri"; styles["Normal"].font.size = Pt(11)
for key, size, color, before, after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)]:
    s=styles[key]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run("SUPPLYPRINT | TECHNICAL PROJECT SUMMARY"), 8.5, True, MUTED)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT; font(footer.add_run("Internal technical brief | 20 June 2026"), 8.5, False, MUTED)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4); r=p.add_run("TECHNICAL PROJECT SUMMARY"); font(r, 24, True, "0B2545")
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16); r=p.add_run("SupplyPrint - physical-product identity, verification, and provenance platform"); font(r, 13, False, MUTED)
for label, value in [("Status", "Local production-like stack healthy"),("Date", "20 June 2026"),("Primary runtime", "React SPA + Spring Boot 3.2 + PostgreSQL/pgvector + ONNX Runtime"),("Database", "PostgreSQL exposed on localhost:5433; service on localhost:10000")]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); font(p.add_run(label + ": "), 10.5, True); font(p.add_run(value), 10.5)

heading(doc, "1. What the project does")
para(doc, "SupplyPrint is an anti-counterfeit and supply-chain provenance system. A manufacturer captures an image of a physical product or package surface. The backend extracts a 128-dimensional visual fingerprint using an ONNX model, stores it in PostgreSQL with pgvector, records a cryptographic feature hash, and queues an immutable ledger attestation through a transactional outbox. A verifier submits a fresh image for the same product identifier; the platform derives a new fingerprint, computes cosine similarity in PostgreSQL, persists a verification audit event, and returns the trusted database/attestation state.")
table(doc, ["Capability", "Current implementation"], [
    ("Enrollment", "JPEG/PNG upload -> server-side ONNX embedding -> PostgreSQL pgvector record -> feature hash -> blockchain outbox"),
    ("Verification", "Fresh JPEG/PNG upload -> ONNX embedding -> single narrow pgvector query -> persisted verification event"),
    ("Evidence", "Product fingerprint, feature hash, enrollment time, transaction metadata, and verification audit events"),
    ("Operations UI", "React command center; metrics and recent activity are retrieved from database-backed APIs"),
    ("Authentication", "Registration, password login, JWT access token, HTTP-only refresh cookie lifecycle"),
], [1.45, 5.05])

heading(doc, "2. Current end-to-end request flow")
heading(doc, "Enrollment", 2)
for text in ["Operator signs in or registers from the SPA.", "The user enters a product identifier, optional metadata, and an actual JPEG or PNG capture.", "POST /api/enroll/image validates type and size (10 MB maximum), runs ONNX inference in the backend, and creates the embedding.", "The enrollment transaction stores product_fingerprints and blockchain_outbox together. This is the transactional outbox pattern: the record is durable before asynchronous ledger delivery.", "The UI shows the real product ID, feature hash, and outbox/ledger state returned by the service."]: bullet(doc, text)
heading(doc, "Verification", 2)
for text in ["A verifier supplies a product ID and a fresh physical image.", "POST /api/verify/image creates an ONNX embedding from the uploaded image; it never accepts a browser-generated vector in this workflow.", "PostgreSQL uses pgvector cosine distance against the stored embedding. The optimized query returns only hash, transaction state, block number, and similarity.", "Every decision is persisted in verification_events, forming the database-backed operations tape and fraud/mismatch count.", "Ledger confirmation is read from persisted attestation state by default, avoiding public-RPC latency in the request path. A strict live-chain mode remains configurable."]: bullet(doc, text)

heading(doc, "3. Architecture and components")
table(doc, ["Layer", "Components and responsibilities"], [
    ("Frontend", "React 18 single-page application. Screens: database-backed dashboard, image enrollment, image verification, evidence vault, authentication."),
    ("API", "Spring Boot product-service with REST controllers for auth, fingerprint enrollment, image verification, evidence lookup, health, metrics, and dashboard."),
    ("AI inference", "OnnxEmbeddingService loads models/fingerprint.onnx at startup and transforms 256x256 grayscale image patches into L2-normalized 128-dimensional vectors."),
    ("Primary data", "PostgreSQL 15 + pgvector: product_fingerprints, blockchain_outbox, verification_events, users, and legacy products."),
    ("Ledger integration", "Web3j + ProductRegistrar contract. Blockchain writes are scheduled from blockchain_outbox and protected by retry/circuit-breaker configuration."),
    ("Observability", "Spring Actuator health, Prometheus endpoint, Grafana/Prometheus Compose services, Micrometer counters/timers, Hikari pool MBeans."),
], [1.35, 5.15])

heading(doc, "4. Key APIs")
table(doc, ["Endpoint", "Purpose", "Persistence"], [
    ("POST /auth/register", "Create operator account and issue session", "users"),
    ("POST /auth/login", "Authenticate operator and refresh-cookie session", "users"),
    ("POST /api/enroll/image", "Enroll image-derived physical fingerprint", "product_fingerprints + blockchain_outbox"),
    ("POST /api/verify/image", "Verify fresh image-derived fingerprint", "product_fingerprints + verification_events"),
    ("GET /api/verify/{productId}/log", "Read provenance/evidence record", "product_fingerprints"),
    ("GET /api/dashboard", "Operational metrics and recent events", "Two optimized PostgreSQL read queries"),
    ("GET /actuator/health", "Service, database, ONNX, and blockchain state", "Runtime health probes"),
], [2.0, 2.8, 1.7])

heading(doc, "5. Database design and data behavior")
para(doc, "The product_fingerprints table is the identity store: a unique product ID, vector(128) embedding, SHA-256 feature hash, optional metadata, transaction hash, block number, and creation time. The blockchain_outbox table ensures the database write and ledger work request are committed atomically. verification_events is an append-only operational/audit record containing product ID, result, similarity, persisted attestation state, and timestamp.")
bullet(doc, "The SPA no longer creates simulated embeddings or dashboard figures.")
bullet(doc, "The Locust scenario discovers product IDs through the service's database dashboard. It does not seed synthetic products or vectors.")
bullet(doc, "When the database is empty, Locust performs real aggregate queries only; provenance reads automatically begin as soon as genuine enrolled product records exist.")

heading(doc, "6. Performance engineering completed")
table(doc, ["Change", "Why it matters"], [
    ("Single verification projection", "Removed a second database lookup and stopped transferring the full embedding row on verification."),
    ("Two-query dashboard read model", "Replaced multiple repository count/top-N calls with one aggregate SQL query plus one bounded recent-events query."),
    ("Focused indexes", "Added pending-outbox dispatch and verification-event time/result indexes."),
    ("Async ledger delivery", "The product verification path reads persisted attestation state instead of waiting for a public blockchain RPC."),
    ("Runtime tuning", "Configured Tomcat worker/accept/connection limits and Hikari MBeans for safe load visibility."),
    ("Deployment fixes", "Moved from Alpine to glibc-compatible JRE for ONNX Runtime; corrected container PORT=10000; aligned UUID schemas with JPA entities."),
], [2.0, 4.5])

heading(doc, "7. Measured live performance")
para(doc, "Benchmark environment: local Docker Compose PostgreSQL/pgvector and product-service, Locust 2.42.6, 30 concurrent users, ramp 5 users/second, 45-second run. The database was empty, so the measured steady-state workload was /api/dashboard PostgreSQL aggregate reads. These are real requests against the running database; they are not browser timings or generated metrics.")
table(doc, ["Metric", "Before optimization", "After optimization"], [
    ("Requests / failures", "430 / 0%", "442 / 0%"),
    ("Average latency", "11 ms", "7 ms"),
    ("p50", "10 ms", "7 ms"),
    ("p95", "18 ms", "11 ms"),
    ("p99", "23 ms", "14 ms"),
], [2.3, 2.1, 2.1])
para(doc, "Interpretation: at the same user count and think-time profile, p95 improved 39% and p99 improved 39%. This benchmark is a low-latency read-path baseline, not a capacity ceiling. A production capacity conclusion requires a representative, approved physical-capture dataset so that image inference, vector verification, audit writes, and ledger queue behavior are all exercised.", "Important: ")

heading(doc, "8. Current live state")
table(doc, ["Item", "State"], [
    ("PostgreSQL", "Healthy; exposed at localhost:5433"),
    ("Product service", "Healthy; exposed at localhost:10000"),
    ("ONNX model", "Loaded successfully in the running container"),
    ("Blockchain", "Disabled in the local Compose environment; outbox entries remain pending until configured"),
    ("Current database records", "0 enrolled physical products and 0 verification events at the benchmark point"),
    ("Locust artifacts", "performance/results/baseline_db_* and performance/results/optimized_db_*"),
], [2.1, 4.4])

heading(doc, "9. Production readiness gaps and next priorities")
para(doc, "The platform is now runnable with real image ingestion and real database persistence, but the following items are required before calling it a production anti-counterfeit system:")
for text in ["Train and validate the ONNX model on a controlled, representative real product-capture dataset. The repository's training script currently demonstrates synthetic-texture training; it is not evidence of field-grade anti-counterfeit accuracy.", "Enable a funded, monitored blockchain RPC plus deployed contract and secret management. The local stack intentionally disables blockchain writes.", "Implement JWT authentication/authorization enforcement on protected API routes; the present security configuration is permissive for development compatibility.", "Introduce tenant ownership on fingerprints and audit records, then enforce tenant/role scoping at every query boundary.", "Add object-storage retention policy and image provenance policy if retaining capture images is required. The present core flow stores derived vectors and metadata, not original photos.", "Run capacity, soak, failure-mode, and image-inference benchmarks with real approved captures and a production-like PostgreSQL volume." ]: bullet(doc, text)

heading(doc, "10. Operating the local stack")
para(doc, "Start the database and service with: docker compose up -d postgresql product-service. Check health at http://localhost:10000/actuator/health. The service expects PostgreSQL internally and publishes port 10000. PostgreSQL maps host port 5433 to container port 5432.")
para(doc, "Run the database-backed benchmark with: python -m locust -f performance/locustfile_supplyprint.py --headless --host http://localhost:10000 -u 30 -r 5 -t 45s --csv performance/results/run_name. The scenario does not manufacture products or feature vectors; it discovers persisted product IDs through the live dashboard.")

doc.save(OUT)
print(OUT)
